import streamlit as st
import os
import pdfplumber
import docx
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import Chroma
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from langchain.docstore.document import Document
from sentence_transformers import SentenceTransformer, util
import google.generativeai as genai
from dotenv import load_dotenv
from typing import List, Dict, Any, Optional, Tuple

# --- Load API Keys ---
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

if not GOOGLE_API_KEY:
    st.error("Google API Key is missing. Please add it to your .env file.")
    st.stop()

try:
    genai.configure(api_key=GOOGLE_API_KEY)
except Exception as e:
    st.error(f"Error configuring Google API: {e}")

# --- Page Configuration ---
st.set_page_config(page_title="AI Document Chat (Google LLM + Rerank)", layout="wide")
st.title(" AI-Powered Document Chatbot (Google LLM + Sentence Rerank)")
st.header(" Upload PDFs, DOCX, TXT, or CSV & Ask Questions", divider="rainbow")

# --- Caching for Text Extraction ---
@st.cache_data
def extract_text_from_file(uploaded_file) -> Optional[str]:
    """Determine file type and extract text accordingly."""
    try:
        file_extension = uploaded_file.name.split(".")[-1].lower()
        content = None
        if file_extension == "pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                content = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        elif file_extension == "docx":
            doc = docx.Document(uploaded_file)
            content = "\n".join([para.text for para in doc.paragraphs])
        elif file_extension == "txt":
            content = uploaded_file.read().decode("utf-8")
        elif file_extension == "csv":
            uploaded_file.seek(0)
            df = pd.read_csv(uploaded_file)
            content = df.to_string()

        if content:
            return '\n'.join(line for line in content.splitlines() if line.strip())
        else:
            return None
    except Exception as e:
        st.error(f"Error extracting text from {uploaded_file.name}: {e}")
        return None

# --- Text Chunking ---
def split_text_into_chunks(text: str, chunk_size: int = 2000, chunk_overlap: int = 200) -> List[str]:
    """Splits extracted text into smaller chunks for embedding."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        add_start_index=True,
    )
    return splitter.split_text(text)

# --- Vector Storage (ChromaDB) ---
@st.cache_resource(show_spinner=" Creating Vector Store...")
def create_vector_store(text_chunks: List[str], file_names: List[str]) -> Optional[Chroma]:
    """Stores text chunks in ChromaDB with metadata (document source)."""
    if not GOOGLE_API_KEY:
        st.error("Google API Key is missing. Cannot create embeddings.")
        return None
    if not text_chunks:
        st.warning("No text chunks found to create vector store.")
        return None

    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        documents = []
        for i, chunk in enumerate(text_chunks):
            metadata = {"source": file_names[i % len(file_names)], "chunk_index": i}
            documents.append(Document(page_content=chunk, metadata=metadata))

        if not documents:
            st.error("No documents could be created from the text chunks.")
            return None

        vector_db = Chroma.from_documents(
            documents=documents,
            embedding=embeddings
        )
        return vector_db
    except Exception as e:
        st.error(f"Error creating vector store: {e}")
        return None

# --- Load Google LLM ---
@st.cache_resource(show_spinner=" Loading Language Model...")
def load_google_llm():
    """Loads the Google Gemini Pro model."""
    try:
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=GOOGLE_API_KEY, temperature=0.2)
        return llm
    except Exception as e:
        st.error(f"Error loading Google gemini-1.5-flash model: {e}")
        return None

# --- Load Sentence Transformer Model for Reranking ---
@st.cache_resource(show_spinner=" Loading Reranking Model...")
def load_reranking_model():
    """Loads the all-mpnet-base-v2 Sentence Transformer model."""
    try:
        model = SentenceTransformer('all-mpnet-base-v2')
        return model
    except Exception as e:
        st.error(f"Error loading Sentence Transformer model: {e}")
        return None

# --- Conversational Chain (Google Gemini) ---
def get_qa_chain_google(llm):
    """Creates a QA pipeline using the loaded Google Gemini model with an improved prompt."""
    prompt_template = """
    You are an AI assistant specialized in answering questions based *only* on the provided context.
    Analyze the context below, which comes from user-uploaded documents.
    Answer the user's question accurately and concisely using *only* the information present in the context.
    If the answer cannot be found within the provided context, explicitly state "The answer is not available in the provided documents."
    Do not make up information or answer based on prior knowledge outside the context.

    Context:
    {context}

    Question:
    {question}

    Answer:
    """
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(llm, chain_type="stuff", prompt=prompt)
    return chain

# --- Handle User Queries with Reranking ---
def process_user_query(user_question: str, vector_db: Chroma, llm, reranker_model) -> Dict[str, Any]:
    """Retrieves relevant context, reranks using Sentence Transformers, and generates a response using the Google LLM."""
    if not vector_db:
        return {"answer": "Vector database is not available.", "sources": []}
    if not llm:
        return {"answer": "Language model is not loaded.", "sources": []}
    if not reranker_model:
        return {"answer": "Reranking model is not loaded.", "sources": []}

    try:
        # 1. Initial Retrieval
        retriever = vector_db.as_retriever(search_kwargs={"k": 10})
        docs = retriever.get_relevant_documents(user_question)

        if not docs:
            return {"answer": "No relevant information found in the documents for your query.", "sources": []}

        # 2. Reranking using Sentence Transformers
        doc_texts = [doc.page_content for doc in docs]
        doc_embeddings = reranker_model.encode(doc_texts)
        query_embedding = reranker_model.encode([user_question])

        # Calculate cosine similarities
        similarities = util.cos_sim(query_embedding, doc_embeddings)[0].tolist()

        # Sort documents by similarity
        ranked_docs = sorted(zip(similarities, docs), key=lambda x: x[0], reverse=True)
        top_n = 3  # Keep top 3 after reranking
        compressed_docs = [doc for score, doc in ranked_docs[:top_n]]

        if not compressed_docs:
            st.warning("Reranker filtered all initial documents. Using top initial document.")
            compressed_docs = docs[:1]
            if not compressed_docs:
                return {"answer": "Could not find relevant documents even after initial search.", "sources": []}

        # 3. Prepare for Generation
        qa_chain = get_qa_chain_google(llm)
        if not qa_chain:
            return {"answer": "Failed to initialize the QA processing chain.", "sources": []}

        # 4. Generation
        response = qa_chain({"input_documents": compressed_docs, "question": user_question})

        # 5. Extract Sources from the RERANKED documents used
        sources = list(set([doc.metadata.get("source", "Unknown Source") for doc in compressed_docs]))

        return {"answer": response["output_text"], "sources": sources}

    except Exception as e:
        st.error(f"Error processing query: {e}")
        return {"answer": f"An error occurred while processing your question: {e}", "sources": []}

# --- Streamlit UI ---
def main():
    """Main function to run the Streamlit application."""

    # --- Sidebar ---
    st.sidebar.title(" Configuration")
    st.sidebar.markdown("Using **Google Gemini Pro** for the Language Model.")
    st.sidebar.markdown("Using **Sentence Transformers** for reranking.")

    st.sidebar.title(" Upload & Process")
    uploaded_files = st.sidebar.file_uploader(
        "Upload Documents (PDF, DOCX, TXT, CSV)", accept_multiple_files=True, type=["pdf", "docx", "txt", "csv"]
    )

    if st.sidebar.button(" Process Uploaded Files", key="process_button"):
        if uploaded_files:
            all_texts = []
            file_names = []
            with st.spinner("Extracting text from files..."):
                for file in uploaded_files:
                    extracted_text = extract_text_from_file(file)
                    if extracted_text:
                        all_texts.append(extracted_text)
                        file_names.append(file.name)
                    else:
                        st.warning(f"Could not extract text from {file.name} or it was empty.")

            if not all_texts:
                st.error("No text could be extracted from the uploaded files.")
            else:
                full_text = "\n\n--- End of Document ---\n\n".join(all_texts)

                with st.spinner("... Splitting text into chunks..."):
                    text_chunks = split_text_into_chunks(full_text)

                vector_db = create_vector_store(text_chunks, list(set(file_names)))

                if vector_db:
                    st.session_state["vector_db"] = vector_db
                    st.success(f" Processed {len(uploaded_files)} files successfully!")
                else:
                    st.error(" Failed to create vector store. Check logs.")
        else:
            st.warning(" Please upload at least one file.")

    st.sidebar.info(" After processing, ask your questions below.")

    # --- Main Chat Area ---
    st.header(" Ask Questions About Your Documents")

    user_question = st.text_input(" Your question:")

    llm = load_google_llm()
    reranker_model = load_reranking_model()

    if user_question:
        if "vector_db" not in st.session_state or st.session_state["vector_db"] is None:
            st.error(" Please upload and process files before asking questions.")
        elif llm is None:
            st.error(" Language model could not be loaded. Check the sidebar for errors.")
        elif reranker_model is None:
            st.error("Reranking model could not be loaded. Check the sidebar for errors.")
        else:
            with st.spinner(" Thinking... Searching documents and generating response..."):
                vector_db = st.session_state["vector_db"]
                result = process_user_query(user_question, vector_db, llm, reranker_model)

                st.markdown("#### Answer:")
                st.write(result["answer"])

                if result["sources"]:
                    st.markdown("---")
                    st.markdown("#### Sources:")
                    for source in sorted(result["sources"]):
                        st.markdown(f"- `{source}`")

# --- Run the App ---
if __name__ == "__main__":
    main()
